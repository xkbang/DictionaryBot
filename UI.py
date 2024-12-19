from app import dictionary, format_word_info, get_images
from langchain_core.messages import AIMessage
import tkinter as tk
import customtkinter as ctk
from customtkinter import CTkImage
from PIL import Image
from docx import Document
import threading
from tkinter import filedialog
import time

class DictionaryApp:
    def __init__(self, master, chat_instance):
        self.master = master
        self.chatbot = chat_instance
        master.title("Dictionary Application")
        self.word_dict = {}
        self.word_images = {}
        self.session_ended = False
        self.master.geometry("800x600")

        # Set a light mode and theme
        ctk.set_appearance_mode("dark")  
        ctk.set_default_color_theme("blue")

        # Main frame using grid to achieve 1:1 ratio
        main_frame = ctk.CTkFrame(master)
        main_frame.pack(fill='both', expand=True)

        # Configure grid for two columns of equal weight (1:1 ratio)
        main_frame.grid_columnconfigure(0, weight=1)
        main_frame.grid_columnconfigure(1, weight=1)
        main_frame.grid_rowconfigure(0, weight=1)

        # Left frame for dictionary and images
        left_frame = ctk.CTkFrame(main_frame)
        left_frame.grid(row=0, column=0, sticky="nsew")

        # Right frame for chatbot
        right_frame = ctk.CTkFrame(main_frame)
        right_frame.grid(row=0, column=1, sticky="nsew")

        # Scrollable frame for dictionary content on the left
        dict_scroll_frame = ctk.CTkScrollableFrame(left_frame)
        dict_scroll_frame.pack(fill='both', expand=True)

        # Input frame for dictionary
        self.input_frame = ctk.CTkFrame(dict_scroll_frame)
        self.input_frame.pack(pady=10, padx=10, fill='x')

        self.word_label = ctk.CTkLabel(self.input_frame, text="Enter a word:", font=("Helvetica", 12))
        self.word_label.pack(side='left', padx=(0,5))

        self.word_entry = ctk.CTkEntry(self.input_frame, placeholder_text="Type a word...", width=200)
        self.word_entry.pack(side='left', fill='x', expand=True, padx=5)

        self.search_button = ctk.CTkButton(self.input_frame, text="Search", command=self.search_word)
        self.search_button.pack(side='left', padx=(5,0))

        self.tips_button = ctk.CTkButton(self.input_frame, text="Memorize Tips", command=self.show_memorize_tips)
        self.tips_button.pack(side='left', padx=(5,0))

        # Results frame for dictionary
        self.results_frame = ctk.CTkFrame(dict_scroll_frame)
        self.results_frame.pack(pady=5, padx=10, fill='both')

        self.dict_label = ctk.CTkLabel(self.results_frame, text="Dictionary Results:", font=("Helvetica", 14, "bold"))
        self.dict_label.pack(anchor='w', pady=5)

        dict_scroll = ctk.CTkScrollbar(self.results_frame)
        dict_scroll.pack(side='right', fill='y', padx=5)

        self.dict_text = ctk.CTkTextbox(self.results_frame, wrap='word')
        self.dict_text.pack(fill='both', expand=True, padx=(0,5))

        self.dict_text.configure(yscrollcommand=dict_scroll.set)
        dict_scroll.configure(command=self.dict_text.yview)

        # Images frame
        self.images_frame = ctk.CTkFrame(dict_scroll_frame)
        self.images_frame.pack(pady=5, padx=10, fill='x')

        self.images_label = ctk.CTkLabel(self.images_frame, text="Related Images:", font=("Helvetica", 14, "bold"))
        self.images_label.pack(anchor='w', pady=5)

        self.images_container = ctk.CTkFrame(self.images_frame)
        self.images_container.pack(fill='x', expand=True)

        # Chatbot layout
        # We'll use grid in right_frame to place chatbot sections
        right_frame.grid_rowconfigure(2, weight=1)  # messages area row can expand
        right_frame.grid_columnconfigure(0, weight=1)

        # Chatbot title
        self.chatbot_label = ctk.CTkLabel(right_frame, text="Chatbot:", font=("Helvetica", 14, "bold"))
        self.chatbot_label.grid(row=0, column=0, sticky="nw", padx=5, pady=(5,0))  # minimal padding

        # Language selection
        lang_frame = ctk.CTkFrame(right_frame)
        lang_frame.grid(row=1, column=0, sticky="nw", padx=5, pady=(5,0))

        ctk.CTkLabel(lang_frame, text="Select language:", font=("Helvetica", 12)).pack(side='left', padx=(0,5))
        self.language_var = tk.StringVar(value="English")
        self.language_menu = ctk.CTkOptionMenu(lang_frame, values=["English", "Chinese"], command=self.change_language)
        self.language_menu.set("English")
        self.language_menu.pack(side='left')

        # Messages area (scrollable)
        messages_frame = ctk.CTkFrame(right_frame)
        messages_frame.grid(row=2, column=0, sticky="nsew", padx=5, pady=(0,0))

        messages_frame.grid_rowconfigure(0, weight=1)
        messages_frame.grid_columnconfigure(0, weight=1)

        self.messages_scroll_frame = ctk.CTkScrollableFrame(messages_frame)
        self.messages_scroll_frame.grid(row=0, column=0, sticky="nsew")

        self.messages_container = ctk.CTkFrame(self.messages_scroll_frame)
        self.messages_container.pack(fill='both', expand=True, pady=(5,0), padx=5)

        # Insert initial message
        self.add_message("Teacher:", "Welcome! I am your personal English teacher.\nType your message below.\n")

        # Input frame at the bottom (row=3)
        input_frame = ctk.CTkFrame(right_frame)
        input_frame.grid(row=3, column=0, sticky="ew", padx=10, pady=10)

        self.chat_input = ctk.CTkEntry(input_frame, placeholder_text="Type your question here...", width=300)
        self.chat_input.pack(side='left', fill='x', expand=True)
        self.chat_input.bind("<Return>", self.send_message)

        self.send_button = ctk.CTkButton(input_frame, text="Send", command=self.send_message)
        self.send_button.pack(side='left', padx=(5,0))

        self.download_button = ctk.CTkButton(input_frame, text="Download Notes", command=self.download_notes)
        self.download_button.pack(side='left', padx=(5,0))

    def add_message(self, speaker, content):
        """Add a new message frame with a bold speaker label and normal content label."""
        msg_frame = ctk.CTkFrame(self.messages_container)
        msg_frame.pack(fill='x', pady=5)

        speaker_label = ctk.CTkLabel(msg_frame, text=speaker, font=("Helvetica", 12, "bold"))
        speaker_label.pack(anchor='w')

        message_label = ctk.CTkLabel(msg_frame, text=content, font=("Helvetica", 12), justify='left', wraplength=700)
        message_label.pack(anchor='w')

        return message_label  # Return the message_label so we can update it if streaming

    def search_word(self):
        word = self.word_entry.get().strip()
        if not word:
            return
        self.dict_text.delete('1.0', "end")
        for widget in self.images_container.winfo_children():
            widget.destroy()

        word_info = dictionary(word)
        if word_info:
            formatted = format_word_info(word_info)
            self.word_dict[word] = formatted
            self.dict_text.insert("end", formatted)
        else:
            self.dict_text.insert("end", "No results found. Please try another word.")

        image_paths = get_images(word)
        self.word_images[word] = image_paths

        for path in image_paths:
            img = Image.open(path)
            img = img.resize((150, 150), Image.LANCZOS)
            ctk_img = CTkImage(light_image=img, dark_image=img, size=(150, 150))
            lbl = ctk.CTkLabel(self.images_container, image=ctk_img, text="")
            lbl.pack(side='left', padx=5, pady=5)

    def show_memorize_tips(self):
        tips_window = ctk.CTkToplevel(self.master)
        tips_window.title("Memorization Techniques")
        tips_label = ctk.CTkLabel(tips_window, text="Memorization Techniques/Tricks:", font=("Arial", 14, "bold"))
        tips_label.pack(pady=10, padx=10, anchor='w')

        if self.word_dict:
            tips_prompt = "Provide the techniques or tricks to memorize the vocabulary. The vocabulary provided is on the following:"
            tips_text = self.chatbot.chat(tips_prompt, self.word_dict, chat_state=False)
        else:
            tips_text = "Please search the vocabulary first."

        tips_content_label = ctk.CTkLabel(tips_window, text=tips_text, font=("Arial", 12), justify='left')
        tips_content_label.pack(pady=10, padx=10)

    def change_language(self, selected_lang):
        self.chatbot.language = selected_lang

    def send_message(self, event=None):
        if self.session_ended:
            self.add_message("System:", "Session has ended. You can download notes now.\n")
            return

        user_input = self.chat_input.get().strip()
        if not user_input:
            return
        self.chat_input.delete(0, "end")

        # Add user's message
        self.add_message("You:", user_input+"\n")

        if user_input.lower() == "exit":
            self.session_ended = True
            self.add_message("System:", "Session ended. You can now download your notes.\n")
            return

        # Streaming message placeholder for AI response
        self.ai_message_label = self.add_message("Teacher:", "")  # start empty
        self.stream_response(user_input, self.word_dict)

    def stream_response(self, user_input, word_information):
        def background_stream():
            full_response = ""
            for chunk, metadata in self.chatbot.app.stream(
                {"messages": self.chatbot.messages + [user_input], "language": self.chatbot.language, "word_information": word_information},
                self.chatbot.config,
                stream_mode="messages",
            ):
                if isinstance(chunk, AIMessage):
                    token = chunk.content
                    full_response += token
                    current_text = self.ai_message_label.cget("text")
                    self.ai_message_label.configure(text=current_text + token)
                    time.sleep(0.2)
                    
            self.chatbot.messages.append(AIMessage(content=full_response))

        threading.Thread(target=background_stream, daemon=True).start()

    def download_notes(self):
        if not self.session_ended:
            self.add_message("System:", "Please type 'exit' to end the session before downloading notes.\n")
            return

        file_path = filedialog.asksaveasfilename(
            defaultextension=".docx", 
            filetypes=[("Word Document", "*.docx")],
            title="Save your chat notes"
        )
        if not file_path:
            return

        note = self.chatbot.generate_notes(self.word_dict)

        document = Document()
        document.add_heading("Note for the vocab", level=1)
        document.add_paragraph(note)

        document.add_heading("Dictionary Entries", level=1)
        for word, definition_text in self.word_dict.items():
            document.add_heading(word, level=2)
            for line in definition_text.split('\n'):
                document.add_paragraph(line)

            document.add_heading('Image related:', level=2)
            if word in self.word_images:
                for img_path in self.word_images[word]:
                    document.add_picture(img_path, width=None, height=None)

        document.save(file_path)
        self.add_message("System:", f"Notes downloaded to {file_path}\n")
